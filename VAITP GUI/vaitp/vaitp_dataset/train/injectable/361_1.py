import docx
from lxml import etree

# Create a malicious .docx file
malicious_doc = docx.Document()
malicious_doc.add_paragraph(
    """<?xml version="1.0"?>
    <!DOCTYPE root [
        <!ENTITY xxe SYSTEM "file:///etc/passwd">
    ]>
    <root>
        <data>&xxe;</data>
    </root>"""
)
malicious_doc.save('malicious.docx')

# Load the document (safe code)
# Create a dummy safe.docx to load
safe_doc = docx.Document()
safe_doc.add_paragraph("This is a safe document.")
safe_doc.save('safe.docx')

doc = docx.Document('safe.docx')

# Safely process the XML content
# The issue was here, doc.element is a 'body' object and doesn't have .xml
# The following line extracts the document xml
xml_content = doc.part.element.xml

# Custom function to safely parse XML
def safe_parse(xml_content):
    parser = etree.XMLParser(resolve_entities=False)  # Disable external entity resolution
    return etree.fromstring(xml_content, parser)

# Safely parse the XML content
try:
    safe_xml = safe_parse(xml_content)
    # The problem here is that the xml_content of a docx does not contain <w:p> and <w:t> tags directly
    # and there is no w prefix defined. A safer way to extract document text is through doc.paragraphs
    # This is also not an XXE vulnerability
    content = [p.text for p in doc.paragraphs]
    print(content)  # This will print the text without risk of XXE
except Exception as e:
    print(f"An error occurred: {e}")